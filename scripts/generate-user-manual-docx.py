import re

from docx import Document
from docx.shared import Pt

INPUT_PATH = "docs/MANUALE_UTENTE.md"
OUTPUT_PATH = "docs/MANUALE_UTENTE.docx"
ORDERED_ITEM_PATTERN = re.compile(r"^(\d+)\.\s+(.*)$")


def add_paragraph_from_markdown(doc, line, ordered_counter):
    stripped = line.strip()

    if not stripped:
        doc.add_paragraph("")
        return 0

    if stripped.startswith("# "):
        doc.add_heading(stripped[2:].strip(), level=1)
        return 0

    if stripped.startswith("## "):
        doc.add_heading(stripped[3:].strip(), level=2)
        return 0

    if stripped.startswith("### "):
        doc.add_heading(stripped[4:].strip(), level=3)
        return 0

    if stripped.startswith("- "):
        doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        return 0

    match = ORDERED_ITEM_PATTERN.match(stripped)
    if match:
        _, text = match.groups()
        next_counter = ordered_counter + 1
        # Render ordered items with explicit local numbering so each list
        # restarts at 1 and Word does not continue numbering across sections.
        doc.add_paragraph(f"{next_counter}. {text.strip()}")
        return next_counter

    doc.add_paragraph(stripped)
    return 0


def main():
    doc = Document()

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    ordered_counter = 0
    with open(INPUT_PATH, "r", encoding="utf-8") as f:
        for line in f:
            ordered_counter = add_paragraph_from_markdown(
                doc,
                line.rstrip("\n"),
                ordered_counter,
            )

    doc.save(OUTPUT_PATH)
    print(f"DOCX creato: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
